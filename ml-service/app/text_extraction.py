"""
app/text_extraction.py

Extracts raw text from uploaded resume files (PDF or DOCX).
Uses pdfplumber for PDFs (better layout-aware extraction than PyPDF2)
and python-docx for Word documents. Falls back to PyMuPDF if
pdfplumber fails on a malformed/scanned PDF.
"""

import io
import re
import pdfplumber
import fitz  # PyMuPDF
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)
        text = "\n".join(text_chunks).strip()
        if text:
            return text
    except Exception:
        pass  # fall through to PyMuPDF fallback

    # Fallback: PyMuPDF (handles some malformed PDFs pdfplumber chokes on)
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]
        # Also pull text from tables (some resumes use table layouts)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def clean_text(raw_text: str) -> str:
    """Normalize whitespace, remove control chars, collapse blank lines."""
    text = re.sub(r"\r\n?", "\n", raw_text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        raw = extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx"):
        raw = extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file type. Only .pdf and .docx are allowed.")
    return clean_text(raw)
